"""Formato estándar para exportación DOCX (escritos judiciales)."""
from __future__ import annotations

from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class DocxExportFormat:
    font_name: str = "Times New Roman"
    font_size_pt: int = 12
    line_spacing: float = 1.5
    space_after_pt: int = 6  # espacio entre párrafos (entremedio)
    margin_cm: float = 3.0  # márgenes izquierdo y derecho simétricos
    margin_top_bottom_cm: float = 2.5
    page_number_format: str = "lowerRoman"  # i, ii, iii…
    alignment: str = "justify"  # left | justify | center | right


DEFAULT_DOCX_FORMAT = DocxExportFormat()

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _set_font(font, fmt: DocxExportFormat) -> None:
    font.name = fmt.font_name
    font.size = Pt(fmt.font_size_pt)


def _ensure_times_new_roman_run(run, fmt: DocxExportFormat) -> None:
    _set_font(run.font, fmt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), fmt.font_name)


def _apply_paragraph_format(paragraph: Paragraph, fmt: DocxExportFormat) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = fmt.line_spacing
    pf.space_after = Pt(fmt.space_after_pt)
    pf.alignment = _ALIGN_MAP.get(fmt.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
    for run in paragraph.runs:
        _ensure_times_new_roman_run(run, fmt)


def _configure_normal_style(doc: Document, fmt: DocxExportFormat) -> None:
    style = doc.styles["Normal"]
    _set_font(style.font, fmt)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = fmt.line_spacing
    pf.space_after = Pt(fmt.space_after_pt)
    pf.alignment = _ALIGN_MAP.get(fmt.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)


def _configure_page_layout(doc: Document, fmt: DocxExportFormat) -> None:
    section = doc.sections[0]
    margin = Cm(fmt.margin_cm)
    section.left_margin = margin
    section.right_margin = margin
    section.top_margin = Cm(fmt.margin_top_bottom_cm)
    section.bottom_margin = Cm(fmt.margin_top_bottom_cm)
    section.gutter = Cm(0)


def _add_page_number_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)


def _configure_roman_page_numbers(doc: Document, fmt: DocxExportFormat) -> None:
    section = doc.sections[0]
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt.page_number_format)
    pg_num_type.set(qn("w:start"), "1")

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(paragraph)


def apply_docx_format(doc: Document, fmt: DocxExportFormat | None = None) -> None:
    """Aplica márgenes, tipografía, interlineado y numeración romana al documento."""
    fmt = fmt or DEFAULT_DOCX_FORMAT
    _configure_page_layout(doc, fmt)
    _configure_normal_style(doc, fmt)
    _configure_roman_page_numbers(doc, fmt)


def add_formatted_paragraph(doc: Document, text: str, fmt: DocxExportFormat | None = None) -> Paragraph:
    fmt = fmt or DEFAULT_DOCX_FORMAT
    paragraph = doc.add_paragraph(text)
    _apply_paragraph_format(paragraph, fmt)
    return paragraph
