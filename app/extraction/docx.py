"""Extracción de texto desde DOCX."""
import io

from docx import Document


def extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("No se pudo extraer texto del documento Word.")
    return text
